# !pip install requests python-docx

import requests
from docx import Document
import os

subscription_key = "API KEY HERE"
endpoint = 'ENDPOINT HERE'
location = "eastus2"
language_destination = 'pt-br'

def translator_text(text, language_destination):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-Key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-Type': 'application/json',
      'X-ClientTraceId': str(os.urandom(16))
  }

  body = [{
      'text': text
  }]
  params = {
      'api-version': 3.0,
      'from': 'en',
      'to': language_destination
  }

  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]

# translator_text("I know you're somewhere out there, somewhere far away", language_destination)

def translate_document(path):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:
      translated_text = translator_text(paragraph.text, language_destination)
      full_text.append(translated_text)

    translated_doc = Document()
    for line in full_text:
      translated_doc.add_paragraph(line)
    path_translated = path.replace(".docx", f"_{language_destination}.docx")
    translated_doc.save(path_translated)
    return path_translated

input_file = "/content/TESTE TRADUÇÃO.docx"
translate_document(input_file)
